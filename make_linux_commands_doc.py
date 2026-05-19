# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "Linux终端常用命令速查.docx"


def set_run_font(run, size=None, bold=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=115, bottom=90, end=115):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, fill=None, font_size=8.6, color="000000", align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    set_paragraph_spacing(p, before=0, after=0, line=1.12)
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        shade_cell(cell, fill)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=15 if level == 1 else 10, after=6, line=1.15)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=15, bold=True, color="1F4D78")
    else:
        set_run_font(run, size=12.5, bold=True, color="2E74B5")
    return paragraph


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    set_paragraph_spacing(paragraph, after=5, line=1.25)
    return paragraph


def add_tip_box(doc, title, items):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [23.8])
    cell = table.cell(0, 0)
    shade_cell(cell, "F2F6FB")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3, line=1.15)
    run = p.add_run(title)
    set_run_font(run, size=10.5, bold=True, color="1F4D78")
    for item in items:
        p = cell.add_paragraph()
        set_paragraph_spacing(p, after=2, line=1.18)
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        run = p.add_run("• ")
        set_run_font(run, size=10, bold=True, color="2E74B5")
        run = p.add_run(item)
        set_run_font(run, size=10)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=6)


def add_command_table(doc, title, commands):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [3.0, 6.2, 7.3, 7.3]
    set_table_geometry(table, widths)
    headers = ["命令", "标准格式", "用途说明", "使用示例"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, fill="E8EEF5", font_size=9.2, color="1F4D78", align="center")
    repeat_table_header(table.rows[0])

    for command, syntax, scene, example in commands:
        cells = table.add_row().cells
        set_cell_text(cells[0], command, bold=True, font_size=8.5, color="1F4D78", align="center")
        set_cell_text(cells[1], syntax, font_size=8.2)
        set_cell_text(cells[2], scene, font_size=8.4)
        set_cell_text(cells[3], example, font_size=8.2)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=5)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Linux 终端常用命令速查")
    set_run_font(run, size=8.5, color="666666")
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=4, after=3, line=1.12)
    run = title.add_run("Linux 终端常用命令速查")
    set_run_font(run, size=24, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=14, line=1.12)
    run = subtitle.add_run("Linux 系统简介、命令标准格式与使用示例")
    set_run_font(run, size=11, color="555555")

    add_heading(doc, "一、Linux 简单介绍", level=1)
    add_body_paragraph(
        doc,
        "Linux 是一种开源、稳定、可定制的操作系统，广泛应用于服务器、云计算、嵌入式设备、网络设备和软件开发环境。"
        "它的核心优势是权限清晰、工具丰富、资源占用低、适合自动化处理，并且拥有大量成熟的命令行工具。"
    )
    add_body_paragraph(
        doc,
        "在学习 Linux 时，终端是最重要的入口。通过终端可以完成目录切换、文件管理、文本查看、软件安装、权限修改、"
        "程序编译、网络连接、进程查看、日志分析和远程传输等任务。掌握常用命令后，后续学习 C 语言、Makefile、"
        "交叉编译、开发板调试和服务器运维都会更高效。"
    )
    add_tip_box(
        doc,
        "学习建议",
        [
            "先记住命令的基本格式，再结合真实路径和文件名反复练习。",
            "看到 [选项] 表示可以省略或按需要添加参数；看到 [文件]、[目录]、[用户] 等表示需要替换成自己的实际内容。",
            "删除、改权限、改所有者、格式化磁盘等命令执行前要先确认路径，避免误操作。",
        ],
    )

    add_heading(doc, "二、常用终端命令列表", level=1)
    command_sections = [
        (
            "1. 目录定位与基础导航",
            [
                ("pwd", "pwd", "显示当前所在目录的完整路径。", "pwd"),
                ("ls", "ls [选项] [目录]", "查看目录内容；-l 显示详细信息，-a 显示隐藏文件。", "ls -la ~/project"),
                ("cd", "cd [目录路径]", "切换工作目录。", "cd ~/project/src"),
                ("cd ..", "cd ..", "返回上一级目录。", "cd .."),
                ("clear", "clear", "清空终端屏幕，便于继续查看输出。", "clear"),
                ("history", "history [数量]", "查看历史命令，方便找回之前执行过的操作。", "history 30"),
                ("tree", "tree [选项] [目录]", "以树形结构查看目录；需要系统已安装 tree。", "tree -L 2 ~/project"),
            ],
        ),
        (
            "2. 文件与目录管理",
            [
                ("mkdir", "mkdir [目录名]", "创建新目录。", "mkdir led_project"),
                ("mkdir -p", "mkdir -p [多级目录]", "一次创建多级目录，父目录不存在时自动创建。", "mkdir -p projects/{led,key,uart}"),
                ("touch", "touch [文件名]", "创建空文件或更新文件修改时间。", "touch main.c Makefile"),
                ("cp", "cp [源文件] [目标位置]", "复制文件。", "cp main.c main_bak.c"),
                ("cp -r", "cp -r [源目录] [目标位置]", "递归复制整个目录。", "cp -r led_project led_project_bak"),
                ("mv", "mv [源路径] [目标路径]", "移动或重命名文件、目录。", "mv old.c main.c"),
                ("rm", "rm [文件名]", "删除文件；执行前确认路径。", "rm main.o"),
                ("rm -r", "rm -r [目录名]", "递归删除目录。", "rm -r old_demo"),
            ],
        ),
        (
            "3. 文件内容查看与统计",
            [
                ("cat", "cat [文件名]", "直接显示小文件内容。", "cat Makefile"),
                ("less", "less [文件名]", "分页查看长文件，可上下翻页和搜索。", "less /var/log/syslog"),
                ("head", "head -n [行数] [文件名]", "查看文件开头若干行。", "head -n 20 main.c"),
                ("tail", "tail -n [行数] [文件名]", "查看文件末尾若干行。", "tail -n 30 build.log"),
                ("tail -f", "tail -f [日志文件]", "持续跟踪文件新增内容，常用于观察日志。", "tail -f run.log"),
                ("wc", "wc [选项] [文件名]", "统计行数、单词数、字节数；-l 统计行数。", "wc -l main.c"),
                ("diff", "diff [文件1] [文件2]", "比较两个文件差异。", "diff old.c main.c"),
            ],
        ),
        (
            "4. 查找、搜索与过滤",
            [
                ("find", "find [路径] [条件]", "按名称、类型、大小等条件查找文件。", "find ~/project -name \"*.c\""),
                ("grep", "grep [选项] \"关键字\" [文件]", "在文件中搜索关键字；-n 显示行号。", "grep -n \"main\" main.c"),
                ("grep -r", "grep -r \"关键字\" [目录]", "递归搜索目录中的文件内容。", "grep -r \"open(\" ~/project"),
                ("which", "which [命令名]", "查看命令实际路径，常用于确认工具是否在 PATH 中。", "which gcc"),
                ("whereis", "whereis [命令名]", "查找命令、源码和手册位置。", "whereis gcc"),
                ("管道 |", "命令1 | 命令2", "把前一个命令输出交给后一个命令处理。", "ls /dev | grep ttyUSB"),
            ],
        ),
        (
            "5. 权限、用户与管理员操作",
            [
                ("chmod", "chmod [权限] [文件]", "修改文件权限。", "chmod +x run.sh"),
                ("chmod 数字", "chmod [三位权限数字] [文件]", "使用数字方式设置权限，例如 755。", "chmod 755 app"),
                ("chown", "chown [用户:用户组] [文件]", "修改文件所有者和用户组。", "sudo chown user:user data.txt"),
                ("sudo", "sudo [命令]", "以管理员权限执行命令。", "sudo apt update"),
                ("whoami", "whoami", "显示当前用户名。", "whoami"),
                ("id", "id [用户名]", "查看用户 UID、GID 和所属用户组。", "id"),
                ("passwd", "passwd [用户名]", "修改用户密码。", "passwd"),
            ],
        ),
        (
            "6. 软件安装与系统服务",
            [
                ("apt update", "sudo apt update", "更新软件包索引，安装软件前建议先执行。", "sudo apt update"),
                ("apt install", "sudo apt install [软件包名]", "安装软件包。", "sudo apt install -y build-essential vim"),
                ("apt remove", "sudo apt remove [软件包名]", "卸载软件包。", "sudo apt remove minicom"),
                ("dpkg -i", "sudo dpkg -i [deb文件]", "安装本地 deb 软件包。", "sudo dpkg -i tool.deb"),
                ("systemctl status", "systemctl status [服务名]", "查看服务运行状态。", "systemctl status ssh"),
                ("systemctl restart", "sudo systemctl restart [服务名]", "重启服务，使配置修改生效。", "sudo systemctl restart ssh"),
            ],
        ),
        (
            "7. 压缩、解压与归档",
            [
                ("tar -czvf", "tar -czvf [压缩包.tar.gz] [文件或目录]", "打包并 gzip 压缩。", "tar -czvf project.tar.gz project"),
                ("tar -xzvf", "tar -xzvf [压缩包.tar.gz]", "解压 tar.gz 文件。", "tar -xzvf project.tar.gz"),
                ("tar -xvf", "tar -xvf [归档文件]", "解包 tar 或 tar.xz 等归档文件。", "tar -xvf gcc-arm.tar.xz"),
                ("zip", "zip -r [压缩包.zip] [目录]", "创建 zip 压缩包。", "zip -r project.zip project"),
                ("unzip", "unzip [压缩包.zip]", "解压 zip 文件。", "unzip docs.zip"),
                ("gzip", "gzip [文件名]", "压缩单个文件生成 .gz 文件。", "gzip build.log"),
                ("gunzip", "gunzip [文件.gz]", "解压 .gz 文件。", "gunzip build.log.gz"),
            ],
        ),
        (
            "8. 编译、构建与程序分析",
            [
                ("gcc", "gcc [源文件.c] -o [输出程序]", "在本机编译 C 程序。", "gcc hello.c -o hello"),
                ("交叉 gcc", "[交叉编译器] [源文件.c] -o [输出程序]", "生成可在目标开发板运行的程序。", "arm-linux-gnueabihf-gcc hello.c -o hello_arm"),
                ("make", "make [目标]", "根据 Makefile 自动构建工程。", "make"),
                ("make clean", "make clean", "执行 Makefile 中的 clean 目标，清理编译产物。", "make clean"),
                ("file", "file [文件名]", "查看文件类型和 CPU 架构。", "file hello_arm"),
                ("ldd", "ldd [可执行文件]", "查看动态库依赖。", "ldd ./hello"),
                ("gdb", "gdb [可执行文件]", "启动调试器定位程序问题。", "gdb ./hello"),
                ("strip", "strip [可执行文件]", "去掉调试符号，减小程序体积。", "strip hello_arm"),
            ],
        ),
        (
            "9. 进程、资源与任务控制",
            [
                ("ps", "ps [选项]", "查看进程；ps aux 常用于查看全部进程。", "ps aux | grep app"),
                ("top", "top", "实时查看 CPU、内存和进程占用。", "top"),
                ("kill", "kill [进程ID]", "结束指定进程。", "kill 12345"),
                ("kill -9", "kill -9 [进程ID]", "强制结束无响应进程，谨慎使用。", "kill -9 12345"),
                ("jobs", "jobs", "查看当前终端启动的后台任务。", "jobs"),
                ("bg", "bg [%任务号]", "让暂停的任务在后台继续运行。", "bg %1"),
                ("fg", "fg [%任务号]", "把后台任务切回前台。", "fg %1"),
                ("Ctrl + C", "快捷键 Ctrl + C", "终止当前前台程序。", "运行程序时按 Ctrl + C"),
            ],
        ),
        (
            "10. 网络、远程登录与文件传输",
            [
                ("ip a", "ip a", "查看网卡和 IP 地址。", "ip a"),
                ("ping", "ping [选项] [IP或域名]", "测试网络连通性。", "ping -c 4 192.168.1.100"),
                ("ssh", "ssh [用户名]@[IP地址]", "远程登录 Linux 主机或开发板。", "ssh root@192.168.1.100"),
                ("scp", "scp [本地文件] [用户]@[IP]:[目标路径]", "把本地文件传到远程设备。", "scp app root@192.168.1.100:/root/"),
                ("scp -r", "scp -r [本地目录] [用户]@[IP]:[目标路径]", "递归传输整个目录。", "scp -r project root@192.168.1.100:/root/"),
                ("wget", "wget [下载链接]", "从网络下载文件。", "wget http://example.com/file.tar.gz"),
                ("curl", "curl [选项] [地址]", "访问网络接口或下载内容。", "curl http://127.0.0.1:8080"),
                ("netstat", "netstat -tunlp", "查看网络连接和监听端口；部分系统需安装 net-tools。", "netstat -tunlp"),
            ],
        ),
        (
            "11. 系统信息、磁盘与设备节点",
            [
                ("uname", "uname [选项]", "查看系统内核、架构和版本信息。", "uname -a"),
                ("lsb_release", "lsb_release -a", "查看 Ubuntu 发行版信息。", "lsb_release -a"),
                ("df", "df -h", "查看磁盘分区使用情况。", "df -h"),
                ("du", "du -sh [路径]", "查看目录占用空间大小。", "du -sh ~/project"),
                ("free", "free -h", "查看内存使用情况。", "free -h"),
                ("dmesg", "dmesg | tail", "查看内核日志，常用于确认 USB 设备识别情况。", "dmesg | tail"),
                ("ls /dev", "ls /dev/[设备模式]", "查看设备节点，例如串口设备。", "ls /dev/ttyUSB*"),
                ("mount", "mount [选项] [设备或共享] [挂载点]", "挂载磁盘、NFS 目录或其他文件系统。", "sudo mount /dev/sdb1 /mnt/usb"),
                ("umount", "umount [挂载点]", "卸载已经挂载的目录。", "sudo umount /mnt/usb"),
            ],
        ),
        (
            "12. 重定向、日志与组合命令",
            [
                (">", "命令 > 文件", "把命令输出写入文件，会覆盖原内容。", "ls -l > file_list.txt"),
                (">>", "命令 >> 文件", "把命令输出追加到文件末尾。", "echo \"test begin\" >> run.log"),
                ("2>", "命令 2> 错误日志", "单独保存错误输出。", "make 2> build_error.log"),
                ("&>", "命令 &> 日志文件", "同时保存标准输出和错误输出。", "make &> build_all.log"),
                ("&&", "命令1 && 命令2", "前一个命令成功后才执行后一个命令。", "make && ./hello"),
                (";", "命令1 ; 命令2", "按顺序执行多个命令，不关心前一个是否成功。", "pwd ; ls -l"),
                ("alias", "alias 名称='命令'", "给常用命令起别名，提高输入效率。", "alias ll='ls -la'"),
            ],
        ),
        (
            "13. 帮助、手册与环境变量",
            [
                ("man", "man [命令名]", "查看命令手册。", "man grep"),
                ("--help", "[命令] --help", "快速查看命令参数说明。", "tar --help"),
                ("whatis", "whatis [命令名]", "用一句话查看命令用途。", "whatis chmod"),
                ("command -v", "command -v [命令名]", "判断命令是否存在，并显示路径。", "command -v gcc"),
                ("echo", "echo [内容]", "输出文本或变量内容。", "echo $PATH"),
                ("export", "export 变量名=变量值", "设置环境变量。", "export PATH=$PATH:/opt/toolchain/bin"),
            ],
        ),
    ]

    for section_title, commands in command_sections:
        add_command_table(doc, section_title, commands)

    add_heading(doc, "三、建议练习顺序", level=1)
    add_tip_box(
        doc,
        "从基础到项目实践",
        [
            "第 1 轮：练习 pwd、ls、cd、mkdir、touch、cp、mv、rm、cat，熟悉路径和文件操作。",
            "第 2 轮：加入 grep、find、chmod、tar、gcc、make，完成普通 C 程序编译。",
            "第 3 轮：加入 ip、ping、ssh、scp、dmesg、mount，开始进行远程连接和设备调试。",
            "第 4 轮：把常用命令组合成项目流程，例如 make && scp && ssh，上板运行并查看日志。",
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
